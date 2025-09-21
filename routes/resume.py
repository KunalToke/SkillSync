import os
import io
from flask import Blueprint, render_template, request, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import json

# Blueprint
resume_bp = Blueprint('resume', __name__)

# Upload folder
UPLOAD_FOLDER = "static/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@resume_bp.route('/resume', methods=['GET', 'POST'])
def resume():
    if request.method == 'POST':
        # Handle photo upload
        photo_file = request.files.get("photo")
        photo_path = None
        if photo_file and photo_file.filename:
            filename = secure_filename(photo_file.filename)
            photo_path = os.path.join(UPLOAD_FOLDER, filename)
            photo_file.save(photo_path)

        # Collect form data
        resume_data = {
            "name": request.form.get("name"),
            "role": request.form.get("role"),
            "email": request.form.get("email"),
            "phone": request.form.get("phone"),
            "address": request.form.get("address"),
            "objective": request.form.get("objective"),
            "skills": request.form.get("skills"),
            "education": request.form.get("education"),
            "experience": request.form.get("experience"),
            "projects": request.form.get("projects"),
            "certifications": request.form.get("certifications"),
            "achievements": request.form.get("achievements"),
            "languages": request.form.get("languages"),
            "hobbies": request.form.get("hobbies"),
            "portfolio": request.form.get("portfolio"),
            "skill_ratings": request.form.get("skill_ratings"),
            "template": request.form.get("template"),
            "references": request.form.get("references"),
            "photo_path": photo_path
        }

        # Resume Scoring
        score, ats_feedback = score_resume(resume_data)

        # Export options
        if "export_pdf" in request.form:
            return export_pdf(resume_data)
        elif "export_docx" in request.form:
            return export_docx(resume_data)
        elif "export_json" in request.form:
            return export_json(resume_data)

        # --- NEW: Render the preview page ---
        return render_template(
            "resume_preview.html",
            resume=resume_data,
            score=score,
            ats_feedback=ats_feedback
        )

    # GET request → Show empty form
    return render_template("resume.html", resume=None)


# ------------ Resume Scoring ------------
def score_resume(data):
    score = 50
    feedback = []

    if data.get("skills"):
        score += 10
    else:
        feedback.append("⚠ No skills added.")

    if data.get("experience"):
        score += 15
    else:
        feedback.append("⚠ Work experience is missing.")

    if data.get("certifications"):
        score += 10

    if data.get("portfolio"):
        score += 15

    return min(score, 100), " | ".join(feedback) if feedback else "✅ Looks good! Resume ATS-friendly."


# ------------ Export Functions ------------
def export_docx(data):
    doc = Document()
    doc.add_heading(data.get("name", "Unnamed Candidate"), 0)
    if data.get("role"):
        doc.add_paragraph(data["role"])

    if data["photo_path"]:
        try:
            doc.add_picture(data["photo_path"], width=Inches(1.5))
        except Exception as e:
            print("Photo error:", e)

    doc.add_paragraph(f"📧 {data.get('email', '')} | 📞 {data.get('phone', '')}")
    if data.get("address"):
        doc.add_paragraph(f"🏠 {data['address']}")

    if data.get("objective"):
        doc.add_heading("Profile", level=1)
        doc.add_paragraph(data["objective"])

    sections = [
        "skills", "education", "experience", "projects",
        "certifications", "achievements", "languages",
        "hobbies", "portfolio", "references"
    ]

    for sec in sections:
        if data.get(sec):
            doc.add_heading(sec.capitalize(), level=1)
            doc.add_paragraph(data[sec])

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name="resume.docx")


def export_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    if data["photo_path"]:
        try:
            story.append(RLImage(data["photo_path"], width=1.5*inch, height=1.5*inch))
        except Exception as e:
            print("PDF Photo error:", e)

    story.append(Paragraph(f"<b><font size=18 color='navy'>{data.get('name', 'Unnamed Candidate')}</font></b>", styles["Title"]))
    if data.get("role"):
        story.append(Paragraph(f"<font size=12>{data['role']}</font>", styles["Normal"]))
    if data.get("objective"):
        story.append(Paragraph(f"<i>{data['objective']}</i>", styles["Normal"]))
    story.append(Paragraph(f"📧 {data.get('email', '')} | 📞 {data.get('phone', '')}", styles["Normal"]))
    if data.get("address"):
        story.append(Paragraph(f"🏠 {data['address']}", styles["Normal"]))
    story.append(Spacer(1, 12))

    for section, value in {
        "Skills": data.get("skills"),
        "Education": data.get("education"),
        "Experience": data.get("experience"),
        "Projects": data.get("projects"),
        "Certifications": data.get("certifications"),
        "Achievements": data.get("achievements"),
        "Languages": data.get("languages"),
        "Hobbies": data.get("hobbies"),
        "Portfolio": data.get("portfolio"),
        "References": data.get("references"),
    }.items():
        if value:
            story.append(Paragraph(f"<b><font color='darkgreen'>{section}</font></b>", styles["Heading2"]))
            story.append(Paragraph(value, styles["Normal"]))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="resume.pdf", mimetype="application/pdf")


def export_json(data):
    json_data = json.dumps(data, indent=4)
    buffer = io.BytesIO(json_data.encode("utf-8"))
    return send_file(buffer, as_attachment=True, download_name="resume.json", mimetype="application/json")
